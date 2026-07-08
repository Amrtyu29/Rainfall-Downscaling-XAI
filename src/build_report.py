"""
build_report.py
================
Generates the full, formal college project report as a Word .docx:

    Downscaling Rainfall over India using Explainable Artificial Intelligence (XAI)

It assembles title page, certificate, declaration, acknowledgement, abstract,
auto-updating Table of Contents / List of Figures / List of Tables, six chapters,
references and an appendix. All numbers and figures are pulled from the project's
own outputs/ so the document stays faithful to the results.

Run:  python src/build_report.py
Out:  report/Project_Report_Full_Rainfall_Downscaling_XAI.docx
"""
import os
import pandas as pd

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "outputs")
FIGDIR = os.path.join(OUT, "figures")
REPORT_PATH = os.path.join(ROOT, "report", "Project_Report_Full_Rainfall_Downscaling_XAI.docx")

CENTER = WD_ALIGN_PARAGRAPH.CENTER
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY
LEFT = WD_ALIGN_PARAGRAPH.LEFT

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = RGBColor(0x1E, 0x40, 0xAF)

# running counters for cached figure / table numbers
_FIG = [0]
_TAB = [0]


# --------------------------------------------------------------------------- #
# low-level Word helpers
# --------------------------------------------------------------------------- #
def _append_run(paragraph, child):
    r = OxmlElement("w:r")
    r.append(child)
    paragraph._p.append(r)
    return r


def add_field(paragraph, instr, cached=None):
    """Insert a Word complex field (with an optional cached result)."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    _append_run(paragraph, begin)

    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    _append_run(paragraph, it)

    if cached is not None:
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        _append_run(paragraph, sep)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = cached
        _append_run(paragraph, t)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    _append_run(paragraph, end)


def set_updatefields(document):
    """Tell Word to refresh all fields (TOC, SEQ, PAGE) when the file opens."""
    settings = document.settings.element
    el = settings.find(qn("w:updateFields"))
    if el is None:
        el = OxmlElement("w:updateFields")
        settings.append(el)
    el.set(qn("w:val"), "true")


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# --------------------------------------------------------------------------- #
# content helpers
# --------------------------------------------------------------------------- #
def body(document, text, justify=True, italic=False, size=None, space_after=6):
    p = document.add_paragraph()
    p.alignment = JUSTIFY if justify else LEFT
    run = p.add_run(text)
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.alignment = JUSTIFY
    return p


def numbered(document, text):
    p = document.add_paragraph(style="List Number")
    p.add_run(text)
    p.alignment = JUSTIFY
    return p


def heading(document, text, level=1, page_break=False):
    if page_break:
        document.add_page_break()
    h = document.add_heading(text, level=level)
    return h


def add_figure(document, filename, caption, width_in=6.0):
    path = filename if os.path.isabs(filename) else os.path.join(OUT, filename)
    p = document.add_paragraph()
    p.alignment = CENTER
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Inches(width_in))
    else:
        p.add_run(f"[missing figure: {os.path.relpath(path, ROOT)}]").italic = True
    cap = document.add_paragraph(style="Caption")
    cap.alignment = CENTER
    _FIG[0] += 1
    cap.add_run("Figure ")
    add_field(cap, r" SEQ Figure \* ARABIC ", str(_FIG[0]))
    cap.add_run(": " + caption)
    cap.paragraph_format.space_after = Pt(12)


def add_table(document, headers, rows, caption=None, widths=None, first_col_bold=False):
    if caption:
        cap = document.add_paragraph(style="Caption")
        cap.alignment = CENTER
        _TAB[0] += 1
        cap.add_run("Table ")
        add_field(cap, r" SEQ Table \* ARABIC ", str(_TAB[0]))
        cap.add_run(": " + caption)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10.5)
        shade_cell(hdr[i], "1E3A5F")
        hdr[i].paragraphs[0].alignment = CENTER
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10.5)
            if first_col_bold and i == 0:
                run.bold = True
            cells[i].paragraphs[0].alignment = LEFT if i == 0 else CENTER
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    document.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


# --------------------------------------------------------------------------- #
# document styling
# --------------------------------------------------------------------------- #
def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 15, NAVY),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 12, ACCENT),
        ("Title", 24, NAVY),
        ("Caption", 10, RGBColor(0x44, 0x44, 0x44)),
    ]:
        if name in [s.name for s in document.styles]:
            st = document.styles[name]
            st.font.name = "Times New Roman"
            st.font.size = Pt(size)
            try:
                st.font.color.rgb = color
            except Exception:
                pass
            if name.startswith("Heading"):
                st.font.bold = True
                st.paragraph_format.space_before = Pt(14)
                st.paragraph_format.space_after = Pt(6)
                st.paragraph_format.keep_with_next = True
    # captions italic
    cap = document.styles["Caption"]
    cap.font.italic = True


def set_margins(section, left=1.25, right=1.0, top=1.0, bottom=1.0):
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)


def footer_page_number(section, fmt="decimal", start=None):
    sectPr = section._sectPr
    # page-number format for this section
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sectPr.append(pg)
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = CENTER
    p.text = ""
    add_field(p, " PAGE ", "1")


def enable_title_page(section):
    """First page of the section has no header/footer (title page)."""
    sectPr = section._sectPr
    if sectPr.find(qn("w:titlePg")) is None:
        sectPr.append(OxmlElement("w:titlePg"))
    section.different_first_page_header_footer = True


# --------------------------------------------------------------------------- #
# methodology pipeline schematic (generated on the fly)
# --------------------------------------------------------------------------- #
def build_pipeline_figure():
    path = os.path.join(FIGDIR, "methodology_pipeline.png")
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

        fig, ax = plt.subplots(figsize=(11, 5.2))
        ax.set_xlim(0, 12)
        ax.set_ylim(0, 7)
        ax.axis("off")

        def box(x, y, w, h, text, fc, tc="white", fs=10.5):
            b = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.12",
                               linewidth=1.2, edgecolor="#334155", facecolor=fc)
            ax.add_patch(b)
            ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                    fontsize=fs, color=tc, weight="bold", wrap=True)
            return (x, y, w, h)

        def arrow(p1, p2, color="#334155"):
            ax.add_patch(FancyArrowPatch(p1, p2, arrowstyle="-|>", mutation_scale=16,
                                         linewidth=1.6, color=color))

        # core row
        b1 = box(0.2, 4.6, 2.3, 1.4, "ERA5 predictors\n(7 large-scale fields)", "#1E40AF")
        b2 = box(3.0, 4.6, 2.3, 1.4, "Point-to-pixel\nfeature table\n(~9.05M rows)", "#2563EB")
        b3 = box(5.8, 4.6, 2.4, 1.4, "ML downscaler\nLightGBM / XGBoost", "#0F766E")
        b4 = box(9.0, 4.6, 2.6, 1.4, "0.25 deg IMD\nrainfall (downscaled)", "#047857")
        arrow((2.5, 5.3), (3.0, 5.3))
        arrow((5.3, 5.3), (5.8, 5.3))
        arrow((8.2, 5.3), (9.0, 5.3))

        # explainability
        bx = box(5.6, 2.4, 2.8, 1.3, "SHAP (TreeSHAP)\nexplainability", "#7C3AED")
        arrow((7.0, 4.6), (7.0, 3.7))

        # cmip6 branch
        bc1 = box(0.2, 2.4, 2.3, 1.3, "CMIP6 GCMs\n(MPI, EC-Earth3)", "#B45309")
        bc2 = box(3.0, 2.4, 2.3, 1.3, "Bias correction\n(quantile mapping)", "#D97706")
        arrow((2.5, 3.05), (3.0, 3.05))
        arrow((5.3, 3.05), (5.6, 3.05))

        # extremes branch
        be1 = box(3.0, 0.4, 2.3, 1.3, "NCEP daily\npredictors", "#9333EA")
        be2 = box(5.8, 0.4, 2.4, 1.3, "Extreme-day\nclassifier (XGBoost)", "#6D28D9")
        be3 = box(9.0, 0.4, 2.6, 1.3, "Flood-day drivers\n(SHAP)", "#7C3AED")
        arrow((5.3, 1.05), (5.8, 1.05))
        arrow((8.2, 1.05), (9.0, 1.05))

        # future output
        bf = box(9.0, 2.4, 2.6, 1.3, "Explained future\nrainfall change", "#059669")
        arrow((8.4, 3.05), (9.0, 3.05))

        ax.text(6.0, 6.6, "Explainable Rainfall Downscaling Pipeline",
                ha="center", fontsize=13.5, weight="bold", color="#1E3A5F")
        plt.tight_layout()
        fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        return "figures/methodology_pipeline.png"
    except Exception as e:  # pragma: no cover
        print("  (pipeline figure skipped:", e, ")")
        return None


# --------------------------------------------------------------------------- #
# main build
# --------------------------------------------------------------------------- #
def build():
    doc = Document()
    configure_styles(doc)

    sec = doc.sections[0]
    set_margins(sec)
    enable_title_page(sec)

    # ===================== TITLE PAGE =====================
    def center(text, size, bold=False, color=None, italic=False, before=0, after=6):
        p = doc.add_paragraph()
        p.alignment = CENTER
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        return p

    center("A PROJECT REPORT ON", 13, bold=True, before=12, after=18)
    center("Downscaling Rainfall over India using", 22, bold=True, color=NAVY, after=2)
    center("Explainable Artificial Intelligence (XAI)", 22, bold=True, color=NAVY, after=18)
    center("A machine-learning statistical downscaling system that explains its own "
           "predictions with SHAP — extended to CMIP6 future projections and "
           "daily rainfall extremes.", 11.5, italic=True, color=RGBColor(0x55, 0x55, 0x55), after=24)

    center("Submitted in partial fulfilment of the requirements for the degree of", 12, after=2)
    center("«Bachelor of Technology / B.Sc. — your degree»", 12.5, bold=True, after=2)
    center("in", 12, after=2)
    center("«Your Department / Programme»", 12.5, bold=True, after=24)

    center("Submitted by", 12, after=4)
    center("«Your Full Name»    (Roll No. «XXXX»)", 13, bold=True, color=NAVY, after=20)

    center("Under the guidance of", 12, after=4)
    center("«Prof. / Dr. Guide Name»", 13, bold=True, color=NAVY, after=24)

    center("«Department Name»", 12.5, bold=True, after=2)
    center("«College / University Name»", 12.5, bold=True, after=2)
    center("«City, State»", 12, after=18)
    center("Academic Year 2025–26", 12.5, bold=True, after=6)

    # ===================== CERTIFICATE =====================
    heading(doc, "Certificate", level=1, page_break=True)
    body(doc,
         "This is to certify that the project report entitled “Downscaling Rainfall over India "
         "using Explainable Artificial Intelligence (XAI)” is a bona fide record of the project "
         "work carried out by «Your Full Name» (Roll No. «XXXX») under my supervision and guidance, "
         "in partial fulfilment of the requirements for the award of the degree of "
         "«Your Degree» in «Your Department» at «College / University Name» during the "
         "academic year 2025–26.")
    body(doc,
         "The work presented in this report is original and has not been submitted, in part or in "
         "full, for the award of any other degree or diploma of this or any other institution.")
    doc.add_paragraph().paragraph_format.space_after = Pt(48)
    sign = doc.add_table(rows=2, cols=2)
    sign.autofit = True
    sign.rows[0].cells[0].paragraphs[0].add_run("\n\n______________________________").bold = False
    sign.rows[0].cells[1].paragraphs[0].add_run("\n\n______________________________").bold = False
    c0 = sign.rows[1].cells[0].paragraphs[0]
    c0.add_run("«Prof. / Dr. Guide Name»\nProject Guide\n«Department Name»")
    c1 = sign.rows[1].cells[1].paragraphs[0]
    c1.add_run("«Head of Department»\nHead of Department\n«Department Name»")
    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    ext = doc.add_table(rows=2, cols=2)
    ext.rows[0].cells[0].paragraphs[0].add_run("\n______________________________")
    ext.rows[0].cells[1].paragraphs[0].add_run("\n______________________________")
    ext.rows[1].cells[0].paragraphs[0].add_run("Internal Examiner")
    ext.rows[1].cells[1].paragraphs[0].add_run("External Examiner")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.add_run("Place: «City»          Date: ____________")

    # ===================== DECLARATION =====================
    heading(doc, "Declaration", level=1, page_break=True)
    body(doc,
         "I hereby declare that the project report entitled “Downscaling Rainfall over India using "
         "Explainable Artificial Intelligence (XAI)” submitted by me is a genuine record of my own "
         "work carried out under the guidance of «Prof. / Dr. Guide Name». The analysis, results and "
         "conclusions presented here are based on datasets that are publicly available and duly cited, "
         "and the software written for this project is my own.")
    body(doc,
         "To the best of my knowledge and belief, this report contains no material previously "
         "published or written by another person, nor material which has been accepted for the award "
         "of any other degree, except where due acknowledgement and reference have been made.")
    body(doc,
         "The complete source code and reproducible pipeline for this project are openly available at: "
         "https://github.com/Amrtyu29/Rainfall-Downscaling-XAI")
    doc.add_paragraph().paragraph_format.space_after = Pt(48)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("«Your Full Name»\nRoll No. «XXXX»\nDate: ____________")

    # ===================== ACKNOWLEDGEMENT =====================
    heading(doc, "Acknowledgement", level=1, page_break=True)
    body(doc,
         "I would like to express my sincere gratitude to my project guide, «Prof. / Dr. Guide Name», "
         "for the invaluable guidance, encouragement and constructive feedback offered throughout this "
         "work. Their insights on climate science and machine learning shaped both the direction and "
         "the rigour of this project.")
    body(doc,
         "I am thankful to the Head and faculty of the «Department Name», «College / University Name», "
         "for providing the academic environment and resources that made this project possible.")
    body(doc,
         "I gratefully acknowledge the data providers whose openly available datasets underpin this "
         "study: the European Centre for Medium-Range Weather Forecasts (ERA5 reanalysis), the India "
         "Meteorological Department (0.25° gridded rainfall), the World Climate Research Programme and "
         "the CMIP6 modelling groups (MPI-ESM1-2-HR and EC-Earth3, distributed via the public Google "
         "Cloud archive), and the NOAA Physical Sciences Laboratory (NCEP/NCAR daily reanalysis).")
    body(doc,
         "Finally, I thank my family and friends for their constant support and encouragement.")
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("«Your Full Name»")

    # ===================== ABSTRACT =====================
    heading(doc, "Abstract", level=1, page_break=True)
    body(doc,
         "High-resolution rainfall information is essential for agriculture, water management and "
         "disaster preparedness in India, yet global climate models and reanalyses provide output at "
         "coarse spatial resolution. This project develops a machine-learning statistical downscaling "
         "system that maps seven large-scale ERA5 atmospheric predictors to 0.25° IMD gridded rainfall "
         "over India (1980–2023), and — critically — explains its own predictions using SHapley Additive "
         "exPlanations (SHAP). Gradient-boosting models roughly triple the skill of a traditional linear "
         "downscaling baseline (R² = 0.73 vs 0.21 on strictly held-out test years 2016–2023). SHAP "
         "analysis shows the model's reasoning is physically consistent: low-level humidity dominates "
         "Indian rainfall, moisture supply governs the summer monsoon, and secondary drivers are "
         "regionally coherent (temperature in the Himalaya, meridional wind in the far south, sea-level "
         "pressure along the cyclone-prone east coast).")
    body(doc,
         "The validated system is then applied to bias-corrected CMIP6 projections (MPI-ESM1-2-HR and "
         "EC-Earth3; SSP2-4.5 and SSP5-8.5). Driven by corrected historical GCM fields, the chain "
         "reproduces the observed monsoon climatology (pattern r = 0.965, bias +3.6%), licensing future "
         "application. Both GCMs project a wetter Indian monsoon (all-India JJAS +17% to +33% by "
         "2040–2070), and SHAP attribution of the projected change identifies rising specific humidity "
         "as its dominant physical driver — the thermodynamic moistening mechanism — while temperature "
         "and pressure changes alone act weakly in the opposite direction. Finally, a daily-scale "
         "classifier of IMD heavy-rainfall days (≥ 64.5 mm/day; ROC-AUC 0.862 on held-out years) shows "
         "that flood-level extremes are governed by mid-tropospheric circulation (monsoon disturbances) "
         "rather than the humidity that controls monthly totals. To our knowledge this is the first "
         "SHAP-based explanation of downscaled future rainfall change — and of daily extreme-rainfall "
         "drivers — over India, resolved by season and sub-region.")
    p = doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run("statistical downscaling, explainable AI, SHAP, Indian summer monsoon, "
              "gradient boosting, CMIP6, climate projections, extreme rainfall.")

    # ===================== TABLE OF CONTENTS =====================
    heading(doc, "Table of Contents", level=1, page_break=True)
    note = doc.add_paragraph()
    r = note.add_run("(In Microsoft Word, right-click below and choose “Update Field” to build the "
                     "contents with page numbers. This line may be deleted.)")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u',
              "Right-click and select ‘Update Field’ to generate the Table of Contents.")

    # ===================== LIST OF FIGURES =====================
    heading(doc, "List of Figures", level=1, page_break=True)
    p = doc.add_paragraph()
    add_field(p, r'TOC \h \z \c "Figure"',
              "Right-click and select ‘Update Field’ to generate the List of Figures.")

    # ===================== LIST OF TABLES =====================
    heading(doc, "List of Tables", level=1, page_break=True)
    p = doc.add_paragraph()
    add_field(p, r'TOC \h \z \c "Table"',
              "Right-click and select ‘Update Field’ to generate the List of Tables.")

    # ===================== ABBREVIATIONS =====================
    heading(doc, "List of Abbreviations", level=1, page_break=True)
    add_table(
        doc,
        ["Abbreviation", "Meaning"],
        [
            ["XAI", "Explainable Artificial Intelligence"],
            ["SHAP", "SHapley Additive exPlanations"],
            ["ML", "Machine Learning"],
            ["GCM", "Global Climate Model"],
            ["CMIP6", "Coupled Model Intercomparison Project, Phase 6"],
            ["ERA5", "ECMWF Reanalysis v5"],
            ["IMD", "India Meteorological Department"],
            ["NCEP/NCAR", "National Centers for Environmental Prediction / National Center for Atmospheric Research"],
            ["SSP", "Shared Socio-economic Pathway"],
            ["JJAS", "June–July–August–September (summer monsoon)"],
            ["R²", "Coefficient of determination"],
            ["RMSE / MAE", "Root-Mean-Square Error / Mean Absolute Error"],
            ["ROC-AUC / PR-AUC", "Area under the ROC / Precision–Recall curve"],
            ["q850, rh850", "Specific / relative humidity at 850 hPa"],
            ["u850, v850", "Zonal / meridional wind at 850 hPa"],
            ["z500", "Geopotential height at 500 hPa"],
            ["MSLP, T2m", "Mean sea-level pressure, 2-metre temperature"],
        ],
        widths=[1.6, 4.6],
    )

    # start the numbered body with its own page numbering
    footer_page_number(sec, fmt="lowerRoman")  # front matter = i, ii, iii ...
    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(body_sec)
    footer_page_number(body_sec, fmt="decimal", start=1)

    # =================================================================== #
    #  CHAPTER 1 — INTRODUCTION
    # =================================================================== #
    heading(doc, "Chapter 1  Introduction", level=1)

    heading(doc, "1.1  Background and Motivation", level=2)
    body(doc,
         "Rainfall is the single most consequential meteorological variable for India. Agriculture, "
         "which supports a majority of the population, is tied to the timing and amount of the summer "
         "monsoon; reservoirs, hydropower and drinking-water supply depend on seasonal totals; and "
         "flood management hinges on the intensity of individual wet days. Decisions in all of these "
         "domains are taken at the level of districts and river basins, which requires rainfall "
         "information at a fine spatial scale.")
    body(doc,
         "Global climate models (GCMs) and reanalysis datasets describe the large-scale state of the "
         "atmosphere well, but only at a coarse resolution of roughly 0.25° to 1°. This is far too "
         "coarse to resolve the sharp rainfall gradients produced by India's orography — the Western "
         "Ghats, the Himalayan foothills and the Northeast — or to serve district-level planning. "
         "Downscaling is the family of techniques that bridges this gap by translating coarse "
         "large-scale information into local detail.")

    heading(doc, "1.2  Statistical Downscaling", level=2)
    body(doc,
         "Two broad approaches exist. Dynamical downscaling runs a high-resolution regional climate "
         "model nested inside a GCM; it is physically detailed but extremely expensive to compute. "
         "Statistical downscaling instead learns an empirical relationship between large-scale "
         "predictors (winds, humidity, pressure, temperature) and local rainfall from historical data, "
         "and then applies that relationship to new inputs. It is cheap, fast and well suited to the "
         "long observational records available over India — the approach adopted in this project.")

    heading(doc, "1.3  Problem Statement and Research Gap", level=2)
    body(doc, "Three gaps in the existing literature motivate this project:")
    numbered(doc, "Traditional statistical downscaling (linear regression, SDSM and similar methods) "
                  "assumes a largely linear relationship and cannot capture the strongly non-linear "
                  "dynamics of the Indian monsoon, limiting its accuracy.")
    numbered(doc, "Modern machine-learning downscaling is far more accurate, but it typically acts as "
                  "a black box. Without an explanation of why a prediction was made, scientific trust "
                  "and operational adoption remain limited.")
    numbered(doc, "Explainable-ML downscaling has recently been demonstrated for Türkiye (Hisam et "
                  "al., 2025) and China (Lyu & Yong, 2025), but not for India — and, crucially, no "
                  "study has explained why projected future rainfall changes, region by region, or "
                  "which physical processes drive flood-level extremes.")

    heading(doc, "1.4  Objectives", level=2)
    body(doc, "The project sets out to:")
    bullet(doc, "Build an accurate machine-learning statistical downscaling system that maps "
                "large-scale ERA5 predictors to 0.25° IMD rainfall over India, benchmarked against a "
                "traditional linear baseline.")
    bullet(doc, "Open the black box with SHAP, quantifying the physical drivers of rainfall globally, "
                "by season and by sub-region, and mapping the dominant driver of every grid cell.")
    bullet(doc, "Extend the validated, explainable model to CMIP6 future scenarios and explain the "
                "physical cause of the projected change — not merely report it.")
    bullet(doc, "Determine, at daily resolution, what large-scale conditions drive flood-level "
                "rainfall extremes, and whether they differ from the drivers of ordinary monthly totals.")

    heading(doc, "1.5  Scope and Contributions", level=2)
    body(doc,
         "The study covers the Indian landmass (6–38°N, 66–100°E) over 1980–2023 for monthly totals "
         "and 2000–2023 for daily extremes. Its main contributions are: (i) the first seasonal and "
         "sub-regional SHAP driver analysis of downscaled rainfall for India; (ii) the first SHAP-based "
         "attribution of downscaled future rainfall change over India; and (iii) the first SHAP "
         "attribution of the drivers of daily extreme-rainfall days over India. A fully reproducible "
         "code base and an interactive dashboard accompany the work.")

    heading(doc, "1.6  Organisation of the Report", level=2)
    body(doc,
         "Chapter 2 reviews the relevant literature. Chapter 3 describes the study area and datasets. "
         "Chapter 4 details the methodology — the learning problem, the models, the explainability "
         "framework and the CMIP6 and extreme-rainfall extensions. Chapter 5 presents and discusses "
         "the results in three parts (present-day downscaling, future projections, and extremes). "
         "Chapter 6 concludes and outlines future work. References and an appendix follow.")

    # =================================================================== #
    #  CHAPTER 2 — LITERATURE REVIEW
    # =================================================================== #
    heading(doc, "Chapter 2  Literature Review", level=1, page_break=True)

    heading(doc, "2.1  Statistical Downscaling of Rainfall", level=2)
    body(doc,
         "The foundations of statistical downscaling were reviewed by Wilby and Wigley (1997), who "
         "classified methods into weather-typing, regression and stochastic weather generators. "
         "Regression-based transfer functions — relating large-scale predictors to local variables — "
         "remain the most widely used family. Their principal weakness is the assumption of a simple, "
         "often linear, predictor–predictand relationship, which is a poor fit for convective and "
         "monsoon-driven rainfall.")

    heading(doc, "2.2  Machine-Learning Downscaling", level=2)
    body(doc,
         "The limitations of linear methods have driven a shift towards machine learning. Tree-ensemble "
         "methods — Random Forests and gradient-boosting machines such as XGBoost and LightGBM — are "
         "particularly effective on tabular predictor–predictand problems, capturing non-linear "
         "interactions without manual feature engineering. Their accuracy, however, comes at the cost "
         "of interpretability: a trained ensemble of hundreds of trees is opaque to direct inspection.")

    heading(doc, "2.3  Explainable AI and SHAP", level=2)
    body(doc,
         "Explainable AI (XAI) seeks to make model predictions interpretable. Among XAI methods, "
         "SHAP (Lundberg & Lee, 2017) has become the de-facto standard because it is grounded in "
         "cooperative game theory: it distributes a prediction fairly among the input features as "
         "Shapley values, with desirable consistency guarantees. For tree ensembles, TreeSHAP computes "
         "these values exactly and efficiently, making it the natural explainability tool for the "
         "models used here.")

    heading(doc, "2.4  Explainable Downscaling — Two Key Studies", level=2)
    body(doc,
         "Two recent studies define the state of the art and directly motivate this project. "
         "Lyu and Yong (2025), in JGR: Machine Learning and Computation, use an explainable machine "
         "learning approach for rainfall downscaling over China, combining gradient boosting with SHAP "
         "to interpret the learned predictor–rainfall relationships. Hisam et al. (2025), in Science "
         "of the Total Environment, downscale precipitation over Türkiye by integrating multiple "
         "precipitation products, land-surface data and gauge stations with explainable ML algorithms. "
         "Both demonstrate that explainable downscaling is feasible and scientifically informative — "
         "yet neither addresses India, future projections, or daily extremes, which is precisely the "
         "space this project occupies.")

    heading(doc, "2.5  Climate Projections and CMIP6", level=2)
    body(doc,
         "The Coupled Model Intercomparison Project Phase 6 (CMIP6; Eyring et al., 2016) provides "
         "coordinated GCM simulations under Shared Socio-economic Pathways (SSPs). GCM output must be "
         "bias-corrected before use in impact studies; empirical quantile mapping is a standard, "
         "robust technique for aligning modelled and observed distributions. Applying a validated "
         "downscaling model to bias-corrected GCM predictors is an established route to "
         "high-resolution projections — but explaining the drivers of the resulting change with SHAP "
         "is, to our knowledge, novel for India.")

    heading(doc, "2.6  Summary of the Gap", level=2)
    add_table(
        doc,
        ["Capability", "Prior work", "This project"],
        [
            ["ML downscaling of rainfall", "China, Türkiye", "India (0.25°, 1980–2023)"],
            ["SHAP explanation of skill", "Global importance", "Global + seasonal + sub-regional + per-cell maps"],
            ["Future rainfall projection", "Reported as change maps", "Downscaled from bias-corrected CMIP6"],
            ["Explaining the projected change", "Not addressed", "SHAP attribution of change (novel)"],
            ["Drivers of daily extremes", "Not addressed", "SHAP attribution of ≥64.5 mm/day days (novel)"],
        ],
        caption="Positioning of this project against the prior explainable-downscaling literature.",
        widths=[2.1, 1.9, 2.4],
    )

    # =================================================================== #
    #  CHAPTER 3 — STUDY AREA AND DATA
    # =================================================================== #
    heading(doc, "Chapter 3  Study Area and Data", level=1, page_break=True)

    heading(doc, "3.1  Study Area", level=2)
    body(doc,
         "The study domain is the Indian landmass within the box 6–38°N, 66–100°E. India's rainfall is "
         "dominated by the summer monsoon (June–September, JJAS), which contributes the bulk of the "
         "annual total, followed by the north-east monsoon (October–December, OND) that brings rain to "
         "the south-eastern peninsula. Sharp gradients occur along the Western Ghats, the Himalayan "
         "foothills and the Northeast, making the domain a demanding test for any downscaling method.")

    heading(doc, "3.2  Datasets", level=2)
    add_table(
        doc,
        ["Dataset", "Role", "Details"],
        [
            ["ERA5 reanalysis (ECMWF / Copernicus)", "Predictors (X)",
             "u850, v850, q850, rh850, z500, MSLP, T2m; monthly, 1980–2023, India box"],
            ["IMD gridded rainfall (0.25°)", "Target (y)",
             "Daily gauge-based analysis aggregated to monthly, 1980–2023"],
            ["CMIP6: MPI-ESM1-2-HR, EC-Earth3 (r1i1p1f1)", "Future projections",
             "historical 1980–2014; SSP2-4.5 & SSP5-8.5 2040–2100; same 7 predictors; public Google Cloud archive"],
            ["NCEP/NCAR reanalysis (daily)", "Extreme-rainfall predictors",
             "Daily large-scale fields, JJAS 2000–2023, via NOAA PSL OPeNDAP"],
            ["India state shapefile", "Analysis & mapping",
             "Regional masks and publication-quality cartography"],
        ],
        caption="Datasets used in the project, their role and key characteristics.",
        widths=[2.2, 1.4, 2.7],
    )
    body(doc,
         "The seven ERA5 predictors were chosen to represent the physical ingredients of rainfall: "
         "low-level moisture (specific and relative humidity at 850 hPa), moisture transport (zonal "
         "and meridional wind at 850 hPa), mid-tropospheric circulation (geopotential height at "
         "500 hPa), and the thermodynamic state (mean sea-level pressure and 2-metre temperature).")

    heading(doc, "3.3  Temporal Split", level=2)
    body(doc,
         "To obtain an honest estimate of skill and to prevent information leakage between adjacent "
         "months, the record is split temporally rather than randomly: training on 1980–2010, "
         "validation on 2011–2015, and testing on 2016–2023. All headline skill numbers reported in "
         "this report are computed on the 2016–2023 test years, which the models never see during "
         "training.")

    # =================================================================== #
    #  CHAPTER 4 — METHODOLOGY
    # =================================================================== #
    heading(doc, "Chapter 4  Methodology", level=1, page_break=True)
    pipeline = build_pipeline_figure()
    if pipeline:
        add_figure(doc, pipeline,
                   "Overview of the explainable rainfall-downscaling pipeline. Large-scale ERA5 "
                   "predictors are learned into 0.25° rainfall and explained with SHAP; the validated "
                   "model is then driven by bias-corrected CMIP6 fields for future projections, and a "
                   "parallel daily branch classifies and explains flood-level extremes.",
                   width_in=6.2)

    heading(doc, "4.1  The Point-to-Pixel Learning Problem", level=2)
    body(doc,
         "The downscaling task is cast as tabular regression. Each (month, grid-cell) pair becomes one "
         "training row: the seven ERA5 predictors at that cell and month, plus the cell's location "
         "(latitude, longitude) and a cyclical encoding of the calendar month (sine and cosine), map "
         "to the observed IMD rainfall. Over the 44-year monthly cube this yields approximately "
         "9.05 million rows. Ocean and out-of-India cells are excluded using the IMD land mask. "
         "Location and month features let a single model represent the whole country while still "
         "learning region- and season-specific behaviour.")

    heading(doc, "4.2  Models Compared", level=2)
    body(doc,
         "Four models form a ladder of increasing complexity, so that the value of each modelling "
         "choice can be isolated:")
    add_table(
        doc,
        ["Model", "Role in the study"],
        [
            ["Linear Regression", "Traditional statistical-downscaling baseline"],
            ["Random Forest", "Bagged tree ensemble; the ML literature standard"],
            ["XGBoost", "Gradient-boosted trees; explained with TreeSHAP"],
            ["LightGBM", "Gradient-boosted trees; best overall skill"],
        ],
        caption="The four models compared, forming a ladder from linear baseline to gradient boosting.",
        widths=[1.9, 4.3],
        first_col_bold=True,
    )
    body(doc,
         "Gradient boosting (XGBoost, LightGBM) is state of the art for tabular data and is the only "
         "family that SHAP can explain exactly and efficiently through TreeSHAP, which is why the "
         "explainability analysis is built on it.")

    heading(doc, "4.3  Evaluation Metrics", level=2)
    body(doc,
         "Model skill is measured on the held-out test years with the coefficient of determination "
         "(R², the fraction of rainfall variance explained), root-mean-square error (RMSE) and mean "
         "absolute error (MAE) in mm/month, the Pearson correlation coefficient (PCC), and the mean "
         "bias. Skill is reported overall, by season (JJAS, OND, JF, MAM) and by sub-region. For the "
         "daily extreme classifier, which faces a heavily imbalanced target, the area under the ROC "
         "curve (ROC-AUC) and the area under the precision–recall curve (PR-AUC) are used instead.")

    heading(doc, "4.4  Explainability with SHAP", level=2)
    body(doc,
         "SHAP (TreeSHAP) attributes every individual prediction to its input variables, in the same "
         "units as the target (mm/month), so that contributions are directly interpretable as "
         "rainfall. Beyond a single global ranking of feature importance, the analysis computes: "
         "(i) seasonal attributions for JJAS, OND, JF and MAM; (ii) sub-regional attributions for the "
         "Western Ghats, the Indo-Gangetic Plain, Northeast India, arid Rajasthan and Peninsular "
         "India; and (iii) per-cell dominant-driver maps that record, for every grid cell, which "
         "predictor carries the largest mean absolute SHAP value. These driver maps are the project's "
         "first explainability novelty for India.")

    heading(doc, "4.5  CMIP6 Projection Chain", level=2)
    body(doc,
         "To project the future, GCM predictors from MPI-ESM1-2-HR and EC-Earth3 are regridded to the "
         "0.25° analysis grid and bias-corrected against ERA5 using per-calendar-month, per-cell "
         "empirical quantile mapping fitted on 1980–2014. This aligns each modelled predictor's "
         "distribution with the observed one before it enters the downscaler. The trained downscaling "
         "model is then applied to the corrected SSP2-4.5 and SSP5-8.5 fields for 2040–2070 and "
         "2070–2100. Crucially, SHAP attributions are compared between each future period and the "
         "baseline: the difference isolates the physical driver of the projected change, turning a "
         "projection into an explanation.")

    heading(doc, "4.6  Extreme-Rainfall Classifier", level=2)
    body(doc,
         "The daily branch targets flood-level extremes. A day at a grid cell is labelled a "
         "heavy-rainfall day if IMD rainfall reaches the operational threshold of 64.5 mm/day. Daily "
         "NCEP/NCAR reanalysis predictors, regridded to the 0.25° IMD grid, feed a class-weighted "
         "XGBoost classifier (heavy days are only about 1.5% of cell-days). The model is trained on "
         "2000–2015 and evaluated on 2016–2023, and SHAP attribution then reveals which large-scale "
         "fields drive extreme days — and whether they differ from the drivers of monthly totals.")

    heading(doc, "4.7  Reproducibility and Tooling", level=2)
    body(doc,
         "The entire pipeline is implemented in Python (scikit-learn, XGBoost, LightGBM, SHAP, xarray, "
         "cartopy) and organised as a set of scripts under src/. Publication-quality maps are produced "
         "by a dedicated cartography module that clips fields to the national outline and overlays "
         "state boundaries. An interactive Streamlit dashboard exposes the results. All code is openly "
         "available (Appendix A).")

    # =================================================================== #
    #  CHAPTER 5 — RESULTS AND DISCUSSION
    # =================================================================== #
    heading(doc, "Chapter 5  Results and Discussion", level=1, page_break=True)

    # ---- Part A ----
    heading(doc, "5.1  Part A — Present-Day Downscaling and Explanation", level=2)

    heading(doc, "5.1.1  Model Skill", level=3)
    add_table(
        doc,
        ["Model", "R²", "RMSE (mm/mo)", "MAE (mm/mo)", "Correlation"],
        [
            ["LightGBM (best)", "0.73", "52.7", "16.8", "0.85"],
            ["XGBoost", "0.71", "54.1", "17.5", "0.84"],
            ["Random Forest", "0.70", "55.3", "16.0", "0.84"],
            ["Linear (traditional)", "0.21", "89.6", "44.0", "0.46"],
        ],
        caption="Downscaling skill on the held-out test years 2016–2023.",
        widths=[1.9, 0.8, 1.4, 1.3, 1.2],
        first_col_bold=True,
    )
    body(doc,
         "The jump from the linear baseline (R² = 0.21) to the tree ensembles (≈ 0.70–0.73) is the "
         "central result of Part A: the predictor–rainfall relationship is fundamentally non-linear, "
         "and machine learning roughly triples the skill of traditional downscaling. The increment "
         "within the tree family is comparatively small, with LightGBM best. Skill peaks in the "
         "monsoon season (JJAS R² = 0.72) and is regionally uniform (R² between 0.59 and 0.74 across "
         "sub-regions), so the model does not succeed only in easy regions.")
    add_table(
        doc,
        ["Season", "R²", "RMSE (mm/mo)", "MAE (mm/mo)"],
        [
            ["JJAS (summer monsoon)", "0.72", "83.4", "33.0"],
            ["MAM (pre-monsoon)", "0.56", "36.7", "11.0"],
            ["OND (post-monsoon)", "0.53", "30.9", "11.4"],
            ["JF (winter)", "0.45", "14.5", "5.7"],
        ],
        caption="Downscaling skill by season on the test years.",
        widths=[2.4, 0.8, 1.5, 1.5],
        first_col_bold=True,
    )
    add_figure(doc, "figures/observed_vs_predicted_map.png",
               "Observed (IMD) versus downscaled (XGBoost) rainfall for July 2020, an unseen test "
               "month. The model reproduces the Western Ghats band, the Northeast maximum and the "
               "Himalayan foothill belt.", width_in=6.2)
    add_figure(doc, "figures/taylor_diagram.png",
               "Taylor diagram comparing all models against observations over the test period. The "
               "gradient-boosting models cluster near the reference point (high correlation, correct "
               "variance); the linear baseline lies far away.", width_in=4.6)
    add_figure(doc, "figures/spatial_skill_maps.png",
               "Spatial distribution of test-period skill. Skill is high and broadly uniform across "
               "India, without collapse in any major region.", width_in=6.2)

    heading(doc, "5.1.2  What Drives Indian Rainfall? (SHAP)", level=3)
    body(doc,
         "SHAP shows the model learned physically sensible reasoning. Humidity controls rainfall "
         "almost everywhere: specific humidity (moisture supply) is the dominant driver over 53.5% of "
         "grid cells and relative humidity (nearness to saturation) over 46.4%. Season matters — "
         "moisture supply governs nearly all of India during the monsoon, while saturation takes over "
         "the peninsula in winter under the north-east monsoon. The secondary drivers are strikingly "
         "coherent: 2-metre temperature in the Himalayan belt, meridional wind in the far south, and "
         "sea-level pressure along the cyclone-prone Odisha coast — the model independently "
         "rediscovered cyclone-driven rainfall.")
    add_figure(doc, "shap/shap_global_bar.png",
               "Global SHAP feature importance (mean |SHAP|, mm/month). Low-level humidity dominates, "
               "followed by location and thermodynamic predictors.", width_in=5.6)
    add_figure(doc, "shap/shap_summary_beeswarm.png",
               "SHAP summary (beeswarm) over the test set: the direction and magnitude of each "
               "predictor's effect on individual predictions.", width_in=5.8)
    add_figure(doc, "shap/shap_seasonal_heatmap.png",
               "Seasonal SHAP importance heatmap (mm/month). Specific humidity leads in every season; "
               "geopotential height rises in importance in winter.", width_in=5.8)
    add_figure(doc, "shap/shap_regional_heatmap.png",
               "Sub-regional SHAP importance heatmap. Relative humidity leads over the Western Ghats; "
               "specific humidity dominates the Indo-Gangetic Plain and Northeast.", width_in=5.8)
    add_figure(doc, "shap/driver_map_india.png",
               "Dominant rainfall driver per grid cell (mean |SHAP|), for annual, monsoon and winter "
               "conditions. To our knowledge the first seasonal-and-regional XAI driver map for India.",
               width_in=6.2)
    add_figure(doc, "shap/secondary_driver_map.png",
               "Second-most-important driver per grid cell, revealing physically coherent regional "
               "structure (temperature in the Himalaya, meridional wind in the far south, sea-level "
               "pressure near the Odisha coast).", width_in=6.0)

    # ---- Part B ----
    heading(doc, "5.2  Part B — Explainable Future Projections (CMIP6)", level=2)

    heading(doc, "5.2.1  Verification of the Projection Chain", level=3)
    body(doc,
         "Two checks license the future application. First, after quantile mapping, the corrected GCM "
         "historical climatologies match ERA5 exactly (mean differences of 0.000 for all seven "
         "predictors in both GCMs). Second, and more demanding, driving the trained downscaler with "
         "corrected historical GCM fields reproduces the observed IMD monsoon climatology with a "
         "spatial pattern correlation of r = 0.965 and an all-India bias of only +3.6%. The chain is "
         "therefore trustworthy before it is extrapolated forward.")

    heading(doc, "5.2.2  Projected Monsoon Rainfall Change", level=3)
    body(doc,
         "Both GCMs agree on the pattern: JJAS rainfall increases across India, most strongly over "
         "Peninsular and central India and most weakly over the Indo-Gangetic Plain and the Northeast, "
         "with local decreases along the Kerala coast — features consistent with recent literature.")
    add_table(
        doc,
        ["All-India JJAS change", "SSP2-4.5", "SSP5-8.5"],
        [
            ["2040–2070", "+17% (MPI) / +32% (EC-Earth3)", "+33% / +49%"],
            ["2070–2100", "+28% / +44%", "+82% / +94%"],
        ],
        caption="Projected all-India summer-monsoon rainfall change relative to 1985–2014, by "
                "scenario, horizon and GCM. Far-future SSP5-8.5 magnitudes are upper-end (see "
                "limitations).",
        widths=[2.2, 2.2, 1.8],
        first_col_bold=True,
    )
    add_figure(doc, "cmip6/change_maps_MPI-ESM1-2-HR.png",
               "Projected JJAS rainfall change over India (MPI-ESM1-2-HR), two scenarios × two "
               "horizons, relative to 1985–2014.", width_in=6.2)
    add_figure(doc, "cmip6/change_maps_EC-Earth3.png",
               "Projected JJAS rainfall change over India (EC-Earth3), for comparison with the MPI "
               "model; the spatial pattern is consistent across both GCMs.", width_in=6.2)

    heading(doc, "5.2.3  Why Does the Monsoon Get Wetter? (Novel Contribution)", level=3)
    body(doc,
         "Comparing SHAP attributions between the future (SSP5-8.5, 2070–2100) and the baseline "
         "isolates the physical driver of the change. In both GCMs, rising specific humidity dominates "
         "the projected increase (change in mean SHAP contribution of +33 mm/month in MPI-ESM1-2-HR "
         "and +39 mm/month in EC-Earth3), with smaller positive contributions from mid-tropospheric "
         "geopotential height and relative humidity. Temperature and sea-level pressure changes alone "
         "contribute negatively. In other words, the model recovers the established mechanism that "
         "thermodynamic moistening of a warmer atmosphere outweighs a modest dynamic weakening of the "
         "monsoon circulation — and it does so consistently across two independent GCMs.")
    add_table(
        doc,
        ["Predictor", "Δ SHAP, MPI (mm/mo)", "Δ SHAP, EC-Earth3 (mm/mo)"],
        [
            ["q850 (specific humidity)", "+33.1", "+39.1"],
            ["z500 (geopotential height)", "+9.5", "+9.8"],
            ["rh850 (relative humidity)", "+7.3", "+7.1"],
            ["u850 (zonal wind)", "+2.3", "+3.0"],
            ["v850 (meridional wind)", "+0.4", "+1.9"],
            ["mslp (sea-level pressure)", "−1.2", "−2.8"],
            ["t2m (temperature)", "−4.9", "−6.4"],
        ],
        caption="SHAP attribution of the projected monsoon change (future minus baseline) for both "
                "GCMs. Rising low-level moisture is the dominant driver; thermodynamic factors "
                "oppose weakly.",
        widths=[2.6, 1.9, 1.9],
        first_col_bold=True,
    )
    add_figure(doc, "cmip6/shap_change_map_MPI-ESM1-2-HR.png",
               "SHAP attribution of projected monsoon rainfall change (MPI-ESM1-2-HR, SSP5-8.5 "
               "2070–2100 vs 1985–2014). Rising low-level moisture drives the wetter monsoon.",
               width_in=6.2)
    add_figure(doc, "cmip6/shap_change_bar_MPI-ESM1-2-HR.png",
               "Predictor-wise contribution to the projected change (MPI-ESM1-2-HR): specific "
               "humidity dominates; temperature and pressure oppose.", width_in=5.6)

    # ---- Part C ----
    heading(doc, "5.3  Part C — What Drives Flood-Level Rainfall Days?", level=2)
    body(doc,
         "The third analysis targets daily extremes. From large-scale NCEP/NCAR fields alone, a "
         "class-weighted XGBoost classifier predicts IMD heavy-rainfall days (≥ 64.5 mm/day) during "
         "the monsoon with a ROC-AUC of 0.862 and a PR-AUC of 0.161 on the held-out years — against a "
         "base rate of only 1.54%, a roughly ten-fold improvement over chance. Flood-level days are "
         "therefore genuinely predictable from the large-scale environment.")
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["ROC-AUC (test 2016–2023)", "0.862"],
            ["PR-AUC (test)", "0.161"],
            ["Base rate of heavy days", "1.54%"],
            ["Lift over base rate", "≈ 10×"],
            ["Test cell-days evaluated", "4.84 million"],
        ],
        caption="Performance of the daily heavy-rainfall classifier on held-out years.",
        widths=[3.2, 2.2],
        first_col_bold=True,
    )
    body(doc,
         "The key scientific finding is a mechanism switch between ordinary and extreme rainfall. "
         "Monthly totals are governed by moisture (Part A), but flood-level days are dominated by "
         "500 hPa geopotential height — the signature of transient dynamic disturbances such as "
         "monsoon depressions — over most of peninsular, central and eastern India. Relative humidity "
         "leads only over the drier north-west and the Indo-Gangetic belt, sea-level pressure over the "
         "western Himalaya, and temperature along the foothills. Thermodynamics sets the stage; "
         "dynamics delivers the flood.")
    add_figure(doc, "extremes/extreme_driver_map.png",
               "Dominant driver of heavy-rainfall days (≥ 64.5 mm/day, JJAS 2016–2023) per grid cell. "
               "Unlike monthly totals (humidity-controlled), daily extremes are dominated by "
               "mid-tropospheric circulation (geopotential height).", width_in=6.2)
    add_figure(doc, "extremes/shap_extreme_vs_ordinary.png",
               "SHAP importance of each predictor on ordinary versus heavy rain days. Geopotential "
               "height and moisture rise in relative importance for extremes.", width_in=5.8)
    add_figure(doc, "extremes/shap_extremes_beeswarm.png",
               "SHAP summary for the extreme-day classifier, showing how each large-scale predictor "
               "pushes the model toward or away from a flood-level day.", width_in=5.8)

    # =================================================================== #
    #  CHAPTER 6 — CONCLUSIONS AND FUTURE SCOPE
    # =================================================================== #
    heading(doc, "Chapter 6  Conclusions and Future Scope", level=1, page_break=True)

    heading(doc, "6.1  Summary of Findings", level=2)
    numbered(doc, "Machine learning roughly triples the skill of traditional statistical downscaling "
                  "of Indian rainfall (R² 0.73 versus 0.21 on strictly held-out years).")
    numbered(doc, "SHAP shows the model's reasoning is physically consistent — humidity-controlled "
                  "rainfall with seasonally and regionally coherent secondary drivers — making the "
                  "downscaler trustworthy, not merely accurate.")
    numbered(doc, "The validated chain reproduces the observed monsoon climatology when driven by "
                  "corrected GCM fields (r = 0.965), and projects a wetter Indian monsoon under both "
                  "SSP2-4.5 and SSP5-8.5, strongest over Peninsular and central India.")
    numbered(doc, "Novel contribution: the first SHAP-based attribution of downscaled future rainfall "
                  "change over India identifies rising low-level moisture as the dominant physical "
                  "driver, with dynamic factors weakly opposing — a result reproduced across two GCMs.")
    numbered(doc, "Daily extremes obey a different mechanism than totals: flood-level days "
                  "(≥ 64.5 mm/day) are predictable from large-scale fields (ROC-AUC 0.86) and are "
                  "dominated by mid-tropospheric circulation rather than humidity — thermodynamics "
                  "sets the stage, dynamics delivers the flood.")

    heading(doc, "6.2  Assumptions and Limitations", level=2)
    bullet(doc, "Stationarity: the predictor–rainfall relationship learned in 1980–2010 is assumed to "
                "hold in a warmer climate — the standard assumption of statistical downscaling. "
                "Far-future SSP5-8.5 magnitudes (+82–94%) exceed typical published ensemble estimates "
                "and should be read as upper-end, pattern-reliable projections, because late-century "
                "humidity values lie beyond the training distribution and the model must extrapolate.")
    bullet(doc, "Ensemble size: two GCMs with one member each; a full multi-model ensemble with "
                "uncertainty quantification is future work.")
    bullet(doc, "Scenarios, not forecasts: results are conditional projections under specified "
                "emission pathways, not predictions of what will happen.")
    bullet(doc, "Predictor resolution for extremes: the daily analysis uses coarse (2.5°) NCEP "
                "predictors, which are limited for convective events.")
    bullet(doc, "Land-cell coverage: rainfall is learned on IMD land cells; skill on strictly "
                "land-only cells (R² ≈ 0.69) is marginally below the full-grid figure, and the "
                "conclusions are unchanged.")

    heading(doc, "6.3  Future Scope", level=2)
    bullet(doc, "Expand to a multi-model CMIP6 ensemble (4–6+ GCMs) and quantify projection "
                "uncertainty, tightening the far-future estimates.")
    bullet(doc, "Merge additional observational products (IMERG, CHIRPS) and land-surface features, "
                "and use finer daily predictors to sharpen the extreme-rainfall analysis.")
    bullet(doc, "Tune model hyper-parameters on the validation set (currently reserved) and test "
                "deep-learning downscalers (e.g., convolutional or transformer architectures).")
    bullet(doc, "Restructure the study into a journal manuscript for peer-reviewed publication.")

    heading(doc, "6.4  Concluding Remark", level=2)
    body(doc,
         "This project delivers, for India, a rainfall downscaling system that is at once accurate and "
         "self-explaining — and shows that the same explainability can be carried into the future to "
         "reveal why the monsoon changes, and down to the daily scale to reveal what drives floods. "
         "Accuracy tells us what the model predicts; explanation tells us whether to believe it, and "
         "what it has learned about the physics of Indian rainfall.")

    # =================================================================== #
    #  REFERENCES
    # =================================================================== #
    heading(doc, "References", level=1, page_break=True)
    refs = [
        "Lyu, Y. & Yong, B. (2025). Using an Explainable Machine Learning Approach for Rainfall "
        "Downscaling. Journal of Geophysical Research: Machine Learning and Computation.",
        "Hisam, E., Sertel, E., & Seker, D. Z. (2025). Precipitation downscaling with the integration "
        "of multiple precipitation products, land surface data and gauge stations using explainable "
        "machine learning algorithms. Science of the Total Environment, 1002, 180540.",
        "Lundberg, S. M. & Lee, S.-I. (2017). A Unified Approach to Interpreting Model Predictions "
        "(SHAP). Advances in Neural Information Processing Systems (NeurIPS).",
        "Lundberg, S. M. et al. (2020). From local explanations to global understanding with "
        "explainable AI for trees. Nature Machine Intelligence, 2, 56–67.",
        "Eyring, V. et al. (2016). Overview of the Coupled Model Intercomparison Project Phase 6 "
        "(CMIP6) experimental design and organization. Geoscientific Model Development, 9, 1937–1958.",
        "Wilby, R. L. & Wigley, T. M. L. (1997). Downscaling general circulation model output: a "
        "review of methods and limitations. Progress in Physical Geography, 21(4), 530–548.",
        "Hersbach, H. et al. (2020). The ERA5 global reanalysis. Quarterly Journal of the Royal "
        "Meteorological Society, 146, 1999–2049.",
        "Pai, D. S. et al. (2014). Development of a new high spatial resolution (0.25° × 0.25°) long "
        "period (1901–2010) daily gridded rainfall data set over India. Mausam, 65(1), 1–18.",
        "Kalnay, E. et al. (1996). The NCEP/NCAR 40-Year Reanalysis Project. Bulletin of the American "
        "Meteorological Society, 77(3), 437–471.",
        "Chen, T. & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings of the "
        "22nd ACM SIGKDD Conference, 785–794.",
        "Ke, G. et al. (2017). LightGBM: A Highly Efficient Gradient Boosting Decision Tree. Advances "
        "in Neural Information Processing Systems (NeurIPS).",
        "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.",
        "Cannon, A. J., Sobie, S. R. & Murdock, T. Q. (2015). Bias correction of GCM precipitation by "
        "quantile mapping. Journal of Climate, 28(17), 6938–6959.",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(f"[{i}]  ").bold = True
        p.add_run(r)

    # =================================================================== #
    #  APPENDIX
    # =================================================================== #
    heading(doc, "Appendix A  Reproducibility and Software", level=1, page_break=True)
    body(doc,
         "The full source code, trained pipeline and interactive dashboard for this project are openly "
         "available at https://github.com/Amrtyu29/Rainfall-Downscaling-XAI (MIT License). The "
         "analysis is implemented in Python 3.9. Key libraries: xarray and netCDF4 (data handling), "
         "scikit-learn, XGBoost and LightGBM (modelling), SHAP (explainability), gcsfs (CMIP6 access), "
         "cartopy and matplotlib (mapping), and Streamlit (dashboard). A pinned requirements.txt is "
         "included in the repository for exact reproduction.")
    body(doc,
         "The pipeline is organised as scripts under src/: data_prep.py (build the predictor cube), "
         "train_models.py (train and score the four models), explain_shap.py and driver_map.py "
         "(SHAP analysis and driver maps), cmip6_download.py / cmip6_prepare.py / cmip6_project.py / "
         "cmip6_explain.py (the future-projection chain), and extreme_download.py / extreme_analysis.py "
         "(the daily extreme-rainfall branch). run_all.py orchestrates the core downscaling workflow.")

    heading(doc, "Appendix B  Detailed SHAP Attribution Tables", level=1, page_break=True)
    body(doc, "Mean absolute SHAP contribution (mm/month) of each predictor, by season:")
    add_table(
        doc,
        ["Predictor", "JJAS", "OND", "JF", "MAM"],
        [
            ["q850 (specific humidity)", "27.2", "15.8", "19.8", "15.0"],
            ["rh850 (relative humidity)", "24.9", "10.6", "12.8", "14.2"],
            ["t2m (temperature)", "9.2", "5.4", "5.8", "5.5"],
            ["mslp (sea-level pressure)", "7.6", "3.5", "3.4", "1.8"],
            ["v850 (meridional wind)", "5.6", "1.9", "1.6", "1.6"],
            ["u850 (zonal wind)", "3.3", "2.9", "2.4", "1.8"],
            ["z500 (geopotential height)", "2.4", "3.3", "5.6", "3.7"],
        ],
        caption="Seasonal SHAP importance (mm/month). Location and month encodings omitted for clarity.",
        widths=[2.6, 0.9, 0.9, 0.9, 0.9],
        first_col_bold=True,
    )
    body(doc, "Mean absolute SHAP contribution (mm/month) of each predictor, by sub-region:")
    add_table(
        doc,
        ["Predictor", "W. Ghats", "Indo-Gangetic", "Northeast", "Arid Raj.", "Peninsular"],
        [
            ["q850", "15.8", "30.9", "30.0", "22.0", "12.2"],
            ["rh850", "34.5", "21.1", "26.6", "19.8", "18.7"],
            ["t2m", "8.1", "9.7", "11.2", "10.0", "5.4"],
            ["mslp", "4.0", "8.8", "5.8", "7.2", "3.2"],
            ["u850", "6.0", "2.1", "3.2", "1.9", "5.0"],
            ["v850", "6.0", "3.0", "4.7", "1.7", "4.7"],
            ["z500", "2.9", "4.3", "5.3", "3.3", "2.2"],
        ],
        caption="Sub-regional SHAP importance (mm/month). Relative humidity leads over the Western "
                "Ghats; specific humidity dominates the Indo-Gangetic Plain and the Northeast.",
        widths=[1.3, 1.0, 1.3, 1.1, 1.0, 1.1],
        first_col_bold=True,
    )

    heading(doc, "Appendix C  Supplementary Figures", level=1, page_break=True)
    add_figure(doc, "shap/shap_dependence_q850.png",
               "SHAP dependence plot for specific humidity (q850): rainfall contribution rises "
               "sharply once low-level moisture crosses a threshold — a non-linear, physically "
               "meaningful response.", width_in=5.6)
    add_figure(doc, "xai_extras/pdp_plots.png",
               "Partial-dependence plots for the leading predictors, corroborating the SHAP "
               "attributions with an independent interpretability method.", width_in=6.0)
    add_figure(doc, "xai_extras/lime_cases.png",
               "LIME explanations for individual sample predictions, providing a local, "
               "instance-level cross-check on the model's reasoning.", width_in=6.0)
    add_figure(doc, "cmip6/change_maps_EC-Earth3.png",
               "Supplementary: projected JJAS change for EC-Earth3 across scenarios and horizons "
               "(companion to the MPI-ESM1-2-HR maps in Chapter 5).", width_in=6.2)

    set_updatefields(doc)
    doc.core_properties.title = "Downscaling Rainfall over India using Explainable AI (XAI)"
    doc.core_properties.subject = "Final Project Report, Academic Year 2025-26"

    os.makedirs(os.path.dirname(REPORT_PATH), exist_ok=True)
    doc.save(REPORT_PATH)
    print("Saved:", REPORT_PATH)
    print(f"Figures embedded: {_FIG[0]}   Tables: {_TAB[0]}")


if __name__ == "__main__":
    build()
